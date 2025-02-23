import streamlit as st
import os
import tempfile
from google.generativeai import configure, GenerativeModel
from langchain.document_loaders import PyPDFLoader, TextLoader
from langchain.vectorstores import FAISS
from langchain.embeddings import HuggingFaceEmbeddings
from dotenv import load_dotenv
from docx import Document
from langchain.schema import Document as LangChainDocument
from docx import Document as DocxDocument
import io
from fpdf import FPDF
import uuid
import sqlite3
import hashlib

st.set_page_config(layout='wide')

load_dotenv()
api_key = os.getenv("GEMINI_API_KEY")

if not api_key:
    st.error("Gemini API key not found. Please set it in the .env file.")
else:
    configure(api_key=api_key)
    model = GenerativeModel("gemini-2.0-flash-exp")

import urllib.parse

def generate_whatsapp_link(chat_history):
    base_url = "https://api.whatsapp.com/send?text="
    message = "\n".join(
        [f"{'You' if chat['role'] == 'user' else 'Bot'}: {chat['message']}" for chat in chat_history]
    )
    encoded_message = urllib.parse.quote(message)
    return f"{base_url}{encoded_message}"


def init_db():
    """Initialize the SQLite database with an additional hash column."""
    conn = sqlite3.connect("uploaded_files.db")
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS files (
                    id INTEGER PRIMARY KEY,
                    filename TEXT NOT NULL,
                    content BLOB NOT NULL,
                    type TEXT NOT NULL,
                    hash TEXT UNIQUE NOT NULL
                )''')
    conn.commit()
    conn.close()

def delete_file_from_db(file_id):
    """Delete a file from the database by its ID."""
    conn = sqlite3.connect("uploaded_files.db")
    c = conn.cursor()
    c.execute("DELETE FROM files WHERE id = ?", (file_id,))
    conn.commit()
    conn.close()
    st.success(f"File with ID {file_id} has been deleted.")

def calculate_file_hash(file):
    """Calculate the hash of a file's content."""
    file_content = file.getvalue()
    return hashlib.sha256(file_content).hexdigest()


def save_file_to_db(file):
    """Save a file to the database if it doesn't already exist."""
    file_hash = calculate_file_hash(file)
    conn = sqlite3.connect("uploaded_files.db")
    c = conn.cursor()

    c.execute("SELECT id FROM files WHERE hash = ?", (file_hash,))
    if c.fetchone():
        st.info(f"The file '{file.name}' is already in the database. Skipping upload.")
    else:
        c.execute("INSERT INTO files (filename, content, type, hash) VALUES (?, ?, ?, ?)",
                  (file.name, file.getvalue(), file.type, file_hash))
        conn.commit()
        st.success(f"File '{file.name}' successfully uploaded and saved to the database.")
    
    conn.close()

init_db()

def get_all_files_from_db():
    conn = sqlite3.connect("uploaded_files.db")
    c = conn.cursor()
    c.execute("SELECT id, filename, type FROM files")
    files = c.fetchall()
    conn.close()
    return files

def get_file_content_from_db(file_id):
    conn = sqlite3.connect("uploaded_files.db")
    c = conn.cursor()
    c.execute("SELECT content, type FROM files WHERE id = ?", (file_id,))
    file = c.fetchone()
    conn.close()
    return file

def load_documents(file):
    """Load and parse documents based on file type."""
    try:
        if file.type == "application/pdf":
            with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_file:
                tmp_file.write(file.getvalue())
                tmp_file_path = tmp_file.name
            loader = PyPDFLoader(tmp_file_path)
            return loader.load()
        elif file.type == "text/plain":
            text = file.getvalue().decode("utf-8")
            return [LangChainDocument(page_content=text)]
        elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
                tmp_file.write(file.getvalue())
                tmp_file_path = tmp_file.name
            docx = DocxDocument(tmp_file_path)
            text = "\n".join([para.text for para in docx.paragraphs])
            return [LangChainDocument(page_content=text)]
        else:
            st.error(f"Unsupported file format: {file.type}")
            return []
    except Exception as e:
        st.error(f"Error processing file {file.name}: {e}")
        return []

def create_faiss_index(documents):
    """Create FAISS index from the loaded documents using HuggingFace embeddings."""
    embeddings = HuggingFaceEmbeddings()
    index = FAISS.from_documents(documents, embeddings)
    return index

def query_gemini(prompt):
    """Send a prompt to the Gemini model and receive a response."""
    try:
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        st.error(f"Error generating response from Gemini model: {e}")
        return ""

def save_chat_as_pdf(chat_history):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    pdf.add_font('Roboto', '', './Roboto-Regular.ttf', uni=True)
    pdf.add_font('Roboto-Bold', 'B', './Roboto-Bold.ttf', uni=True)
    pdf.set_font('Roboto', '', 12)

    pdf.cell(200, 10, txt="Chat History", ln=True, align="C")
    pdf.ln(10)

    for chat in chat_history:
        role = "You" if chat['role'] == 'user' else "Bot"
        pdf.multi_cell(0, 10, txt=f"{role}: {chat['message']}")
        pdf.ln(2)

    pdf_output = io.BytesIO()
    pdf_data = pdf.output(dest='S')
    pdf_output.write(pdf_data)
    pdf_output.seek(0)
    return pdf_output

def create_docx(content):
    """Generate a DOCX file from the given content."""
    doc = DocxDocument()
    for chat in content:
        line = f"{'You' if chat['role'] == 'user' else 'Bot'}: {chat['message']}"
        doc.add_paragraph(line)
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer

col_left,_,col_right = st.columns([1,6,1])

with col_left:
    st.image("Altibbe logo dark.png", width=150)

with col_right:
    st.image("Hedamo.jpg", width=200)

st.title("Report Generator and ChatBot")
st.write("Upload files (PDF, TXT, DOCX) and interact with the chatbot.")

uploaded_files = st.file_uploader("Upload Files", type=["pdf", "txt", "docx"], accept_multiple_files=True)

source_option = st.radio("Select source for chat interaction", ["Uploaded Files", "Files from Database"])

chat_using_uploaded_files=source_option

if uploaded_files:
    for file in uploaded_files:
        save_file_to_db(file)
    st.success("Files successfully uploaded and saved to the database.")

if chat_using_uploaded_files=="Uploaded Files":
    documents = []
    for file in uploaded_files:
        file_documents = load_documents(file)
        if file_documents:
            documents.extend(file_documents)

    if documents:
        index = create_faiss_index(documents)
        st.success("Documents successfully loaded and indexed.")

        if 'chat_history' not in st.session_state:
            st.session_state.chat_history = []

        st.markdown("### Chat History")
        for chat in st.session_state.chat_history:
            if chat['role'] == 'user':
                st.markdown(f"**You**: {chat['message']}")
            elif chat['role'] == 'bot':
                st.markdown(f"**Bot**: {chat['message']}")

        user_query = st.text_area("Ask a question based on the uploaded files:", height=100)

        if st.button("Send Query"):
            st.session_state.chat_history.append({"role": "user", "message": user_query})

            retriever = index.as_retriever()
            relevant_docs = retriever.get_relevant_documents(user_query)

            if relevant_docs:
                context = "\n\n".join(set(doc.page_content for doc in relevant_docs))
                prompt_text = (
    "Perform a high-level AI-assisted compliance screening, evaluating specific metrics such as:\n"
    "1. **Category Fit**: Ensuring the product falls under an eligible category (e.g., food, cosmetics, clothing).\n"
    "2. **NOVA Classification Alignment**: For food products, ensuring they fall within NOVA I-III categories.\n"
    "3. **Organic Certification Verification**: Confirming the presence of valid certifications (e.g., USDA Organic, EU Organic) for all categories.\n"
    "4. **Transparency Metrics**: Assessing the percentage of end-to-end transparency provided by the client across sourcing, processing, and supply chain levels.\n"
    "5. **Compliance Likelihood**: Using proprietary algorithms to flag potential non-compliance, incomplete data, or missing certifications.\n\n"
    
    "**Classify the product into one of four categories:**\n"
    "1. **Platinum Label Eligibility**: 100% end-to-end transparency across all sourcing and production.\n"
    "2. **Gold Label Eligibility (Tier 1)**: Products demonstrating 100% transparency at the client’s level for 75% of the total production process. "
    "The remaining 25%, including upstream inputs and processes, must comply with strict organic certification standards.\n"
    "3. **Green Label Eligibility**: Products demonstrating 100% transparency at the client’s level for 50% of the total production process. "
    "The remaining 50%, including upstream inputs and external processes, must comply with strict organic certification standards.\n"
    "4. **Non-Compliant**: Product does not meet basic criteria (e.g., ultra-processed, lacks necessary organic certification, or falls outside eligible categories).\n\n"

    "**Share a Preliminary Feedback Report, including:**\n"
    "- **Eligibility status**\n"
    "- **Label tier classification**\n"
    "- **Next steps** (audit, advisory, or improvement recommendations)\n\n"

    "**Introduce Category-Specific Advisory Services:**\n"
    "- **Cosmetics**: Assess ingredient safety, sourcing transparency, and compliance with eco-friendly standards.\n"
    "- **Clothing**: Evaluate sourcing of raw materials (organic fibers, ethical labor practices) and traceability of the supply chain.\n"
    "- **Future Categories**: Maintain flexibility to add additional verticals (e.g., wellness products, ethical tech), aligning each with Hedamo’s transparency framework."
)
                prompt = f"Context: {context}\n\nUser: {user_query}+{prompt_text}\nBot: Give the reference links of external sources used for verification. You have access to internet . Please provide links when asked. Also give detailed analysis of the questions asked."

                response = query_gemini(prompt)

                st.session_state.chat_history.append({"role": "bot", "message": response})

                st.markdown("### Latest Interaction")
                st.markdown(f"**You**: {user_query}")
                st.markdown(f"**Bot**: {response}")

                latest_interaction = [{"role": "user", "message": user_query}, {"role": "bot", "message": response}]

                chat_history_content = "\n\n".join(
                    f"{'You' if chat['role'] == 'user' else 'Bot'}: {chat['message']}" for chat in st.session_state.chat_history
                )

                st.markdown("### Download Options")

                st.download_button(
                    label="Download Latest Interaction as PDF",
                    data=save_chat_as_pdf(latest_interaction), 
                    file_name="latest_interaction.pdf",
                    mime="application/pdf"
                )
                st.download_button(
                    label="Download Latest Interaction as DOCX",
                    data=create_docx(latest_interaction),
                    file_name="latest_interaction.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

                st.download_button(
                    label="Download Chat History as PDF",
                    data=save_chat_as_pdf(st.session_state.chat_history), 
                    file_name="chat_history.pdf",
                    mime="application/pdf"
                )

                st.download_button(
                    label="Download Chat History as DOCX",
                    data=create_docx(st.session_state.chat_history),
                    file_name="chat_history.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
                import urllib.parse

                def generate_whatsapp_link(content):
                    """Create a WhatsApp link to share the chat history."""
                    base_url = "https://api.whatsapp.com/send?text="
                    message = "\n".join(
                    [f"{'You' if chat['role'] == 'user' else 'Bot'}: {chat['message']}" for chat in content]
                    )
                    encoded_message = urllib.parse.quote(message)
                    return f"{base_url}{encoded_message}"



                st.markdown("### Share Chat History")

                if st.session_state.chat_history:
                   whatsapp_link = generate_whatsapp_link(st.session_state.chat_history)

                   st.markdown(
                f'<a href="{whatsapp_link}" target="_blank" style="text-decoration: none;">'
               '<button style="background-color:#25D366; color:white; border:none; padding:10px; font-size:16px; cursor:pointer;">'
                'Share via WhatsApp</button></a>',
                unsafe_allow_html=True,
                )
                   
                else:
                    st.info("No chat history to share yet.")
                
            else:
                st.error("No relevant documents found. Try asking another question.")
    else:
        st.error("No valid documents found. Please upload supported file types.")


elif chat_using_uploaded_files == "Files from Database":

    st.markdown("### Select Files for Processing")
    files = get_all_files_from_db()

    selected_file_ids = []

    if files:
        selected_file_ids = st.multiselect("Select files to process:", options=[(file[0], file[1]) for file in files], format_func=lambda x: x[1])
    else:
        st.warning("Please upload files to continue...")

    to_delete=True
    
    if st.button("Delete Selected Files"):
     if selected_file_ids:
        for file_id, _ in selected_file_ids:  
            delete_file_from_db(file_id)  
        st.success("Selected files deleted successfully!")
        to_delete=False
     else:
        st.warning("No files selected.")


    elif selected_file_ids and to_delete:
        all_documents = []
        for file_id in selected_file_ids:
            file_content, file_type = get_file_content_from_db(file_id[0])
            file_like_object = io.BytesIO(file_content)
            file_like_object.name = file_id[1]
            file_like_object.type = file_type

            documents = load_documents(file_like_object)
            all_documents.extend(documents)

        if all_documents:
            index = create_faiss_index(all_documents)
            st.success("Files successfully loaded and indexed.")

            if 'chat_history' not in st.session_state:
                st.session_state.chat_history = []

            st.markdown("### Chat History")
            for chat in st.session_state.chat_history:
                if chat['role'] == 'user':
                    st.markdown(f"**You**: {chat['message']}")
                elif chat['role'] == 'bot':
                    st.markdown(f"**Bot**: {chat['message']}")

            user_query = st.text_area("Ask a question based on the uploaded files:", height=100)

            if st.button("Send Query"):
                st.session_state.chat_history.append({"role": "user", "message": user_query})
                retriever = index.as_retriever()
                relevant_docs = retriever.get_relevant_documents(user_query)

                if relevant_docs:
                    context = "\n\n".join(set(doc.page_content for doc in relevant_docs))
                    prompt_text = (
    "Perform a high-level AI-assisted compliance screening, evaluating specific metrics such as:\n"
    "1. **Category Fit**: Ensuring the product falls under an eligible category (e.g., food, cosmetics, clothing).\n"
    "2. **NOVA Classification Alignment**: For food products, ensuring they fall within NOVA I-III categories.\n"
    "3. **Organic Certification Verification**: Confirming the presence of valid certifications (e.g., USDA Organic, EU Organic) for all categories.\n"
    "4. **Transparency Metrics**: Assessing the percentage of end-to-end transparency provided by the client across sourcing, processing, and supply chain levels.\n"
    "5. **Compliance Likelihood**: Using proprietary algorithms to flag potential non-compliance, incomplete data, or missing certifications.\n\n"
    
    "**Classify the product into one of four categories:**\n"
    "1. **Platinum Label Eligibility**: 100% end-to-end transparency across all sourcing and production.\n"
    "2. **Gold Label Eligibility (Tier 1)**: Products demonstrating 100% transparency at the client’s level for 75% of the total production process. "
    "The remaining 25%, including upstream inputs and processes, must comply with strict organic certification standards.\n"
    "3. **Green Label Eligibility**: Products demonstrating 100% transparency at the client’s level for 50% of the total production process. "
    "The remaining 50%, including upstream inputs and external processes, must comply with strict organic certification standards.\n"
    "4. **Non-Compliant**: Product does not meet basic criteria (e.g., ultra-processed, lacks necessary organic certification, or falls outside eligible categories).\n\n"

    "**Share a Preliminary Feedback Report, including:**\n"
    "- **Eligibility status**\n"
    "- **Label tier classification**\n"
    "- **Next steps** (audit, advisory, or improvement recommendations)\n\n"

    "**Introduce Category-Specific Advisory Services:**\n"
    "- **Cosmetics**: Assess ingredient safety, sourcing transparency, and compliance with eco-friendly standards.\n"
    "- **Clothing**: Evaluate sourcing of raw materials (organic fibers, ethical labor practices) and traceability of the supply chain.\n"
    "- **Future Categories**: Maintain flexibility to add additional verticals (e.g., wellness products, ethical tech), aligning each with Hedamo’s transparency framework."
)

                    prompt = f"Context: {context}\n\nUser: {user_query} +{prompt_text}\nBot: Give the reference links of external sources used for verification. You have access to internet. Please provide links when asked. Give detailed analysis of questions asked."
                    st.write(prompt)
                    response = query_gemini(prompt)

                    st.session_state.chat_history.append({"role": "bot", "message": response})

                    st.markdown("### Latest Interaction")
                    st.markdown(f"**You**: {user_query}")
                    st.markdown(f"**Bot**: {response}")

                    latest_interaction = [{"role": "user", "message": user_query}, {"role": "bot", "message": response}]

                    st.markdown("### Download Options")

                    st.download_button(
                        label="Download Latest Interaction as PDF",
                        data=save_chat_as_pdf(latest_interaction),
                        file_name="latest_interaction.pdf",
                        mime="application/pdf"
                    )
                    st.download_button(
                        label="Download Latest Interaction as DOCX",
                        data=create_docx(latest_interaction),
                        file_name="latest_interaction.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                    st.download_button(
                        label="Download Chat History as PDF",
                        data=save_chat_as_pdf(st.session_state.chat_history),
                        file_name="chat_history.pdf",
                        mime="application/pdf"
                    )

                    st.download_button(
                        label="Download Chat History as DOCX",
                        data=create_docx(st.session_state.chat_history),
                        file_name="chat_history.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

                    whatsapp_link = generate_whatsapp_link(st.session_state.chat_history)

                    st.markdown(
                        f'<a href="{whatsapp_link}" target="_blank" style="text-decoration: none;">'
                        '<button style="background-color:#25D366; color:white; border:none; padding:10px; font-size:16px; cursor:pointer;">'
                        'Share via WhatsApp</button></a>',
                        unsafe_allow_html=True,
                    )
                else:
                    st.error("No relevant documents found. Try asking another question.")
    else:
        st.warning("Please select at least one file to enable the chatbot.")